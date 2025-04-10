import re
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from utils.common_utils import get_llm_api_url, get_llm_api_headers, success, failure


def submit_to_llm(prompt, config):
    system_prompt = """
        You cannot chat with user, just do: Rewrite the given text in a better format,
        ensuring clarity, coherence, and professional tone. Maintain the original meaning.
        Use # for headings, ## for level 2 headings, and ### for level 3 headings.
    """

    conversation = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': prompt},
    ]

    url = get_llm_api_url(config)
    headers = get_llm_api_headers(config)
    payload = {'messages': conversation, 'temperature': 0.6}

    response = requests.post(url, json=payload, headers=headers)

    if response.status_code == 200:
        data = response.json()
        return success(data['choices'][0]['message']['content'])
    else:
        error_message = f'Error: {response.status_code}, {response.text}'
        return failure(code=response.status_code, message=error_message)


def format_to_docx(text):
    doc = Document()
    lines = text.split('\n')

    for line in lines:
        line = clean_text(line)

        if line.startswith('# '):
            title_paragraph = doc.add_heading(line[2:], level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('## '):
            title_paragraph = doc.add_heading(line[3:], level=2)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('### '):
            new_paragraph = doc.add_heading(line[4:], level=3)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('#### '):
            new_paragraph = doc.add_heading(line[5:], level=3)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.italic = True

        elif re.match(r'^\d+\. \*\*(.*?)\*\*', line):
            match = re.match(r'^(\d+)\. \*\*(.*?)\*\*', line)
            paragraph = doc.add_paragraph(f'{match.group(1)}. ', style='List Number')
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('**') and line.endswith('**'):
            new_paragraph = doc.add_paragraph()
            run = new_paragraph.add_run(line[2:-2])
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        else:
            new_paragraph = doc.add_paragraph(line)
            new_paragraph.paragraph_format.first_line_indent = Pt(24)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def clean_text(text):
    return text.strip()
