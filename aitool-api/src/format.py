import requests
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import re
import os

def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def submit_to_llm(prompt):

    system_prompt = """
    You cannot chat with user, just do: Rewrite the given text in a better format,
    ensuring clarity, coherence, and professional tone. Maintain the original meaning.
    Use # for headings, ## for level 2 headings, and ### for level 3 headings.
    """

    conversation = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": prompt}
    ]

    url = f"{basicUrl}/deployments/{modelName}/chat/completions/?api-version={apiVersion}"
    headers = {'Content-Type': 'application/json', 'api-key': apiKey}
    payload = {'messages': conversation, 'temperature': 0.6}

    response = requests.post(url, json=payload, headers=headers)

    if response.status_code == 200:
        data = response.json()
        return data['choices'][0]['message']['content']
    else:
        return f"Error: {response.status_code}, {response.text}"

def clean_text(text):
    return text.strip()

def format_and_save_docx(content, output_path):

    new_doc = Document()
    lines = content.split('\n')

    for line in lines:
        line = clean_text(line)

        if line.startswith('# '): 
            title_paragraph = new_doc.add_heading(line[2:], level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('## '):
            title_paragraph = new_doc.add_heading(line[3:], level=2)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('### '):  
            new_paragraph = new_doc.add_heading(line[4:], level=3)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('#### '):  
            new_paragraph = new_doc.add_heading(line[4:], level=3)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.italic = True
                
        elif re.match(r'^\d+\. \*\*(.*?)\*\*', line): 
            match = re.match(r'^(\d+)\. \*\*(.*?)\*\*', line)
            paragraph = new_doc.add_paragraph(f"{match.group(1)}. ", style='List Number')
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('**') and line.endswith('**'): 
            new_paragraph = new_doc.add_paragraph()
            run = new_paragraph.add_run(line[2:-2])
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        else: 
            new_paragraph = new_doc.add_paragraph(line)
            new_paragraph.paragraph_format.first_line_indent = Pt(24)
            for run in new_paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)

    new_doc.save(output_path)
    print(f"Formatted document saved as: {output_path}")

def process_file(input_txt, output_docx):
    text_content = read_txt_file(input_txt)
    improved_text = submit_to_llm(text_content)
    format_and_save_docx(improved_text, output_docx)

def run_file_converter():
    input_txt_file = input("Enter the path of the TXT file: ")
    if not input_txt_file.lower().endswith('.txt'):
        print("Error: Please provide a .txt file.")
        return

    directory = os.path.dirname(input_txt_file)
    base_name = os.path.basename(input_txt_file).replace(".txt", "")
    output_docx_file = os.path.join(directory, f"{base_name}_converted.docx")

    process_file(input_txt_file, output_docx_file)
